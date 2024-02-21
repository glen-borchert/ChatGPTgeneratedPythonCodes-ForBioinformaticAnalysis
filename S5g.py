import os
from collections import defaultdict, OrderedDict
from multiprocessing import Pool
from docx import Document

def get_unique_20mers(filename):
    sequences = OrderedDict()
    with open(filename, 'r') as file:
        for line in file:
            line = line.strip()
            if not line.startswith('>'):
                for i in range(len(line) - 19):
                    sequence = line[i:i + 20]
                    if sequence not in sequences:
                        sequences[sequence] = 0
    return sequences

def count_sequences(args):
    sequence, second_file = args
    with open(second_file, 'r') as file:
        count = 0
        for line in file:
            line = line.strip()
            count += line.count(sequence)
        return sequence, count

if __name__ == "__main__":
    first_file = input("Enter the path to the first file: ")
    second_file = input("Enter the path to the second file: ")
    output_folder = r"C:\Python"
    output_file = os.path.join(output_folder, "max9.docx")

    unique_sequences = get_unique_20mers(first_file)

    with Pool() as pool:
        sequence_counts = OrderedDict(pool.map(count_sequences, [(seq, second_file) for seq in unique_sequences.keys()]))

    doc = Document()
    doc.add_heading('Sequence Counts', 0)
    for sequence, count in sequence_counts.items():
        doc.add_paragraph(f'{sequence}\t{count}')

    doc.save(output_file)
    print(f"Results saved to {output_file}")
