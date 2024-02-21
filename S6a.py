from Bio import SeqIO
from docx import Document
import os

def count_reads(fasta_file):
    count = 0
    with open(fasta_file, "r") as handle:
        for line in handle:
            if line.startswith(">"):
                count += 1
    return count

def save_to_word_file(output_filename, reads_count, input_filename):
    doc = Document()
    doc.add_paragraph(input_filename)
    doc.add_paragraph(f"Reads count: {reads_count}")
    doc.save(output_filename)

if __name__ == "__main__":
    try:
        fasta_file_path = input("Enter the path to the genomic fasta file: ")
        output_filename = input("Enter the name for the output Word file (without extension): ")

        reads_count = count_reads(fasta_file_path)
        print(f"The number of reads in the file is: {reads_count}")

        file_name = os.path.splitext(os.path.basename(fasta_file_path))[0]
        save_to_word_file(f"{output_filename}.docx", reads_count, file_name)
        print(f"Output saved to {output_filename}.docx")
    except FileNotFoundError:
        print("File not found. Please provide a valid file path.")
    except Exception as e:
        print(f"An error occurred: {e}")
